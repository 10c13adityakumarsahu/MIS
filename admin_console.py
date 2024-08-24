from tkinter import *
import random
from tkinter import ttk
import mysql.connector
from tkinter import messagebox as tmsg
#from admin_login import syslog 
from PIL import ImageTk, Image
from tkinter import filedialog
from tkcalendar import Calendar,DateEntry
from tkPDFViewer import tkPDFViewer as pdf
import os
import docx
from docx.shared import Inches
from multiprocessing import Process, Queue
import cv2
from PIL import *

#=========================Initialization=======================
Win=Tk()
Win.title("Adminisatrator Access")
Win.geometry("720x720")
backcolour3="#7A2048"
style=("timesnewroman",20,"bold")
style1=("Calibri", 20, "italic")
style2=("Calibri", 15, "italic")
style3=("Calibri", 25, "italic")
front="white"
B="#d0d9d9"
backcolour2="black"
Buttoncolor="#7A2048"
buttontxt="white"
backcolour="#408EC6"
backcolour4="#4a9ad4"

l1=[]
#===========Admin User id===================================
def selected(a,b="sec"):
    global user
    global l1
    user=a
    l1.append(user)
    c=b
    sql_commands(user,c)

def sql_commands(a,b="select"):
    if a=="Student":
        tab="students"
        dab="Student_Info"
    elif a=="Instructor":
        tab="instructor"
        dab="Teacher_Info"
    else:
        dab="Admin_Info"
        tab="admin_info"
    mydb = mysql.connector.connect(
  host="localhost",
  user="root",
  password="Aditya@1234",
  database=f"{dab}")
    mycursor = mydb.cursor()

    if b=="add":
        try:
            sql = f"INSERT INTO {tab} VALUES (%s, %s,%s,%s)"
            val = (Val1_input.get(), Val2_input.get(),Val3_input.get(),1)
            mycursor.execute(sql, val)
            mydb.commit()
            mycursor.execute(f"Select * from {tab}")
            passkey=mycursor.fetchall()
            tmsg.showinfo("Success","Task executed")
            reset()
        except:
            tmsg.showerror("Duplicate Entry","User Already Exists or Data Invalid")
            reset()
    if b=="revoke":
        try:
            if tab=="students":
                val = Val1_input.get()
                sql = f"update {tab} set access=0 where registration={val}"
                mycursor.execute(sql)
                mydb.commit()
                mycursor.execute(f"Select * from {tab}")
                passkey=mycursor.fetchall()
                #for x in passkey:
                    #print(x)
                reset()
                tmsg.showinfo("Success","Acess Denied")
            elif tab=="instructor":
                vall=(Val1_input.get())
                sql = f"update {tab} set access=0 where Emp_id={vall}"
                mycursor.execute(sql)
                mydb.commit()
                mycursor.execute(f"Select * from {tab}")
                passkey=mycursor.fetchall()
                #for x in passkey:
                    #print(x)
                tmsg.showinfo("Success","User Data Updated")
                reset()
            else:
                vall=(Val1_input.get())
                sql = f"update {tab} set access=0 where admin_id={vall}"
                mycursor.execute(sql)
                mydb.commit()
                mycursor.execute(f"Select * from {tab}")
                passkey=mycursor.fetchall()
                #for x in passkey:
                    #print(x)
                tmsg.showinfo("Success","User Data Updated")
                reset()
        except:
            tmsg.showerror("Error","Unable To Update Access")
            reset()
    if b=="access":
        try:
            if tab=="students":
                vall=(Val1_input.get())
                sql = f"update {tab} set access=1 where Registration={vall}"
                mycursor.execute(sql)
                mydb.commit()
                mycursor.execute(f"Select * from {tab}")
                passkey=mycursor.fetchall()
                #for x in passkey:
                    #print(x)
                tmsg.showinfo("Success","User Data Updated")
                reset()
            elif tab=="instructor":
                vall=(Val1_input.get())
                sql = f"update {tab} set access=1 where Emp_id={vall}"
                mycursor.execute(sql)
                mydb.commit()
                mycursor.execute(f"Select * from {tab}")
                passkey=mycursor.fetchall()
                #for x in passkey:
                    #print(x)
                tmsg.showinfo("Success","User Access Updated")
                reset()
            else:
                vall=(Val1_input.get())
                sql = f"update {tab} set access=1 where admin_id={vall}"
                mycursor.execute(sql)
                mydb.commit()
                mycursor.execute(f"Select * from {tab}")
                passkey=mycursor.fetchall()
                #for x in passkey:
                    #print(x)
                tmsg.showinfo("Success","User Access Updated")
                reset()

        except:
            tmsg.showerror("Error","Unable To Update Access")
            reset()
    if b=="delete":
        try:
            if tab=="students":
                vall=(Val1_input.get())
                sql = f"delete from {tab}  where Registration={vall}"
                mycursor.execute(sql)
                mydb.commit()
                mycursor.execute(f"Select * from {tab}")
                passkey=mycursor.fetchall()
                #for x in passkey:
                    #print(x)
                tmsg.showinfo("Success","User Deleted")
                reset()
            elif tab=="instructor":
                vall=(Val1_input.get())
                sql = f"delete from {tab} where Emp_id={vall}"
                mycursor.execute(sql)
                mydb.commit()
                mycursor.execute(f"Select * from {tab}")
                passkey=mycursor.fetchall()
                #for x in passkey:
                    #print(x)
                tmsg.showinfo("Success","User Deleted")
                reset()
            else:
                vall=(Val1_input.get())
                sql = f"Delete from {tab} where admin_id={vall}"
                mycursor.execute(sql)
                mydb.commit()
                mycursor.execute(f"Select * from {tab}")
                passkey=mycursor.fetchall()
                #for x in passkey:
                    #print(x)
                tmsg.showinfo("Success","User Deleted")
                reset()

        except:
            tmsg.showerror("Error","Unable To Update Access")
            reset()
    if b=="update":
        try:
            if tab=="students":
                vall=(Val2_input.get(),Val3_input.get(),Val1_input.get())
                sql = f"update {tab} set name=%s,password=%s where Registration=%s"
                mycursor.execute(sql,vall)
                mydb.commit()
                mycursor.execute(f"Select * from {tab}")
                passkey=mycursor.fetchall()
                #for x in passkey:
                    #print(x)
                reset()
                tmsg.showinfo("Success","User Data Updated")
                reset()
            elif tab=="instructor":
                vall=(Val2_input.get(),Val3_input.get(),Val1_input.get())
                sql = f"update {tab} set name=%s,password=%s where Emp_id=%s"
                mycursor.execute(sql,vall)
                mydb.commit()
                mycursor.execute(f"Select * from {tab}")
                passkey=mycursor.fetchall()
                #for x in passkey:
                    #print(x)
                reset()
                tmsg.showinfo("Success","User Data Updated")
                reset()
            else:
                vall=(Val2_input.get(),Val3_input.get(),Val1_input.get())
                sql = f"update {tab} set name=%s,password=%s where Emp_id=%s"
                mycursor.execute(sql,vall)
                mydb.commit()
                mycursor.execute(f"Select * from {tab}")
                passkey=mycursor.fetchall()
                #for x in passkey:
                    #print(x)
                reset()
                tmsg.showinfo("Success","User Data Updated")
                reset()

        except:
            reset()
    if b=="data":
        try:
            if tab=="students":
                reset_1()
                mycursor.execute(f"Select Registration,Name,password,access from {tab}")
                passkey=mycursor.fetchall()
                t1.config(state="normal")
                t1.delete(1.0,END)
                for x in passkey:
                    #print(x)
                    for j in x:
                        t1.insert(END,f"[{j}]\t")
                    t1.insert(END,f"\n")
                t1.config(state="disabled")
            elif tab=="instructor":
                reset_1()
                mycursor.execute(f"Select Emp_id,Name,password,access from {tab}")
                passkey=mycursor.fetchall()
                t1.config(state="normal")
                t1.delete(1.0,END)
                for x in passkey:
                    #print(x)
                    for j in x:
                        t1.insert(END,f"[{j}]\t")
                    t1.insert(END,f"\n")
                t1.config(state="disabled")
            else:
                reset_1()
                mycursor.execute(f"Select * from {tab}")
                passkey=mycursor.fetchall()
                t1.config(state="normal")
                t1.delete(1.0,END)
                for x in passkey:
                    #print(x)
                    for j in x:
                        t1.insert(END,f"[{j}]\t")
                    t1.insert(END,f"\n")
                t1.config(state="disabled")

        except:
            tmsg.showerror("Error","unable to fetch Data")

var=StringVar()
var1=IntVar()
varx=StringVar()
#============Captcha=========================================
def back_command():
    if tmsg.askyesno("Back","Do you want to go back \n unsaved changes will not reflect")==True:
        Win.destroy()
        import admin
    else:
        pass
captcha=random.randint(10000,99999)
def generate_new_captcha():
    global captcha
    captcha1=random.randint(10000,99999)
    Val8.config(text=captcha1)
    Val8_input.delete(0,END)
    captcha=captcha1
def reset():
    Val1_input.delete(0,END)
    Val2_input.delete(0,END)
    Val3_input.delete(0,END)
    Val4.deselect()
    Val8_input.delete(0,END)
    generate_new_captcha()
def validatior():
    if var1.get()==0:
            tmsg.showerror("Error","Accept To Continue")
    if str(captcha)!=Val8_input.get():
            tmsg.showerror("Error Captcha","Wrong Captcha! Retry")
            reset()
def add_user():
    a=var.get()
    validatior()
    selected(a,b="add")
def revoke_user():
    a=var.get()
    validatior()
    selected(a,b="revoke")
def allow_access():
    a=var.get()
    validatior()
    selected(a,b="access")
def delete_user():
    a=var.get()
    validatior()
    selected(a,b="delete")

def change_data():
    a=var.get()
    validatior()
    selected(a,b="update")

def get_data():
    a=var.get()
    selected(a,b="data")

def reset_1():
    global t1
    t1.config(state="normal")
    t1.delete(1.0,END)

def generate_cert():
    import generate_certificate

def open_image():
    file_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg")])
    if file_path:
        xwin=Toplevel(Win)
        img = Image.open(file_path)
        img = img.resize((300, 300), Image.LANCZOS) # Resize image if needed
        img = ImageTk.PhotoImage(img)
        label = Label(xwin, image=img)
        label.image = img # Keep reference to avoid garbage collection
        label.pack()
def stu_files():
    try:
        file_path = filedialog.askopenfilename(initialdir=f"C:/Users/Adity/Desktop/mis/Files/Student/{int(Val1_input.get())}",filetypes=[("Image Files", "*.png *.jpg *.jpeg")])
        if file_path:
            xwin=Toplevel(Win)
            img = Image.open(file_path)
            img = img.resize((300, 300), Image.LANCZOS) # Resize image if needed
            img = ImageTk.PhotoImage(img)
            label = Label(xwin, image=img)
            label.image = img # Keep reference to avoid garbage collection
            label.pack()  
    except:
        tmsg.showerror("Error","Enter Valid Registration Number") 

def instruct_files():
    try:
        file_path = filedialog.askopenfilename(initialdir=f"C:/Users/Adity/Desktop/mis/Files/Instructor/{int(Val1_input.get())}",filetypes=[("Image Files", "*.png *.jpg *.jpeg")])
        if file_path:
            xwin=Toplevel(Win)
            img = Image.open(file_path)
            img = img.resize((300, 300), Image.LANCZOS) # Resize image if needed
            img = ImageTk.PhotoImage(img)
            label = Label(xwin, image=img)
            label.image = img # Keep reference to avoid garbage collection
            label.pack()  
    except:
        tmsg.showerror("Error","Enter Valid Employee Id")    
    
    
def show_det_inst():
    mydb = mysql.connector.connect(
    host="localhost",
    user="root",
    password="Aditya@1234",
    database="Teacher_info")
    mycursor = mydb.cursor()
    detins= Val1_input.get()
    if detins=='':
        tmsg.showerror("Error","Unable To Fetch Data")
        return
    try:
        mycursor.execute(f"SELECT * FROM instructor where Emp_id={detins}")
        passkey = mycursor.fetchall()
        t1.config(state="normal")
        t1.delete(1.0,END)
        m=["Emp_id:","Name:","Password:","Access:","Age:","Number:","Gender:","Email:","Address:","Department:","Qualification:","Salary:","DOJ:","DOE:","A/C No:"]
        count=0
        for x in passkey:
            for y in x:
        #print(x)
                t1.insert(END,f"{m[count]}{y}\t")
                t1.insert(END,f"\n")
                count+=1
        t1.config(state="disabled")
    except:
        tmsg.showerror("Error","Unable To Fetch Data")
def show_fees():
    pass
def upload_file():
    file_path = filedialog.askopenfilename(defaultextension="C:/Users/Adity/Downloads/",filetypes=[("Image Files", "*.png *.jpg *.jpeg")])
    if file_path:
        # Process the file (e.g., save it to a new location)
        save_file(file_path)
def save_file(file_path):
    save_path = filedialog.asksaveasfilename(initialdir="C:/Users/Adity/Desktop/mis/Certificates_uploaded",filetypes=[("Image Files", "*.png *.jpg *.jpeg")])
    if save_path:
        with open(file_path, "rb") as f:
            data = f.read()
        with open(save_path, "wb") as f:
            f.write(data)
def generate_Admission_no():
    global adminssion_no
    adminssion_no=random.randint(100000,1000000)
    complete_Admission(adminssion_no)

def complete_Admission(a):
    mydb = mysql.connector.connect(
    host="localhost",
    user="root",
    password="Aditya@1234",
    database="Student_info")
    mycursor = mydb.cursor()
    acval=int(0)
    val=(a,ad3a.get(),a,0,ad4v.get(),ad5v.get(),ad6v.get(),ad7a.get(),ad8a.get(),ad11a.get(),ad12a.get(),ad13a.get(),120000,0,120000,str(cal.get()),ad10a.get())
    st=f"Insert into students values (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
    mycursor.execute(st,val)
    mydb.commit()
    mycursor.close()
    tmsg.showinfo("Admission",f"Admission Done! \n Registration number is : {a}")
def make_admission():
    global ad3a,ad4v,ad5v,ad6v,cal,ad7a,ad8a,ad10a,ad11a,ad12a,ad13a
    Admission=Toplevel(Win)
    Admission.geometry("800x800")
    Admission.state("zoomed")
    Admission.maxsize(800,800)
    Admission.minsize(800,800)
    Admission.title("New Admission")
    #generate_Admission_no()
    #print(scs)
    ad1=Label(Admission,text="New Admission",background="black",fg="white",font=style)
    ad1.place(relx=0,rely=0,relwidth=1,relheight=0.09)
    ad2=Frame(Admission,bg=backcolour)
    ad2.place(rely=0.1,relx=0,relwidth=1,relheight=0.9)
    ad3=Label(ad2,text="Name",width=20,anchor="w",bg=backcolour4,font=style3)
    ad3.grid(row=0,column=0,pady=2,padx=2)
    ad3a=Entry(ad2,textvariable=StringVar(),width=13,bg="white",font=style3)
    ad3a.grid(row=0,column=1,padx=10)
    ad4=Label(ad2,text="Gender",width=20,anchor="w",bg=backcolour4,font=style3)
    ad4.grid(row=1,column=0,pady=2,padx=2)
    ad4v=StringVar()
    ad4a=OptionMenu(ad2,ad4v,"Male","Female")
    ad4v.set("Select")
    ad4a.grid(row=1,column=1)
    ad5=Label(ad2,text="Admission Branch",width=20,anchor="w",bg=backcolour4,font=style3)
    ad5.grid(row=2,column=0,pady=2,padx=2)
    ad5v=StringVar()
    ad5a=OptionMenu(ad2,ad5v,"Computer Science","Information Technology","Civil Engineering","Mechanical Engineering","Chemical Engineering","Artificial Intelligence","Electronics and Communication","Mining Engineering")
    ad5a.grid(row=2,column=1)
    ad5v.set("Select")
    ad6=Label(ad2,text="Category",width=20,anchor="w",bg=backcolour4,font=style3)
    ad6.grid(row=3,column=0,pady=2,padx=2)
    ad6v=StringVar()
    ad6a=OptionMenu(ad2,ad6v,"Select","General","OBC(NCL)","SC","ST","PWD")
    ad6a.grid(row=3,column=1)
    ad6v.set("Select")
    ad7=Label(ad2,text="Mother's Name",width=20,anchor="w",bg=backcolour4,font=style3)
    ad7.grid(row=4,column=0,pady=2,padx=2)
    ad7a=Entry(ad2,textvariable=StringVar(),width=13,bg="white",font=style3)
    ad7a.grid(row=4,column=1,padx=10)
    ad8=Label(ad2,text="Father's Name",width=20,anchor="w",bg=backcolour4,font=style3)
    ad8.grid(row=5,column=0,pady=2,padx=2)
    ad8a=Entry(ad2,textvariable=StringVar(),width=13,bg="white",font=style3)
    ad8a.grid(row=5,column=1,padx=10)
    ad9=Label(ad2,text="Date Of Birth",width=20,anchor="w",bg=backcolour4,font=style3)
    ad9.grid(row=6,column=0,pady=2,padx=2)
    sel=StringVar() # declaring string variable 
    cal=DateEntry(ad2,selectmode='day',textvariable=sel)
    cal.grid(row=6,column=1,padx=20)
    ad10=Label(ad2,text="Mother Tounge",width=20,anchor="w",bg=backcolour4,font=style3)
    ad10.grid(row=7,column=0,pady=2,padx=2)
    ad10a=Entry(ad2,textvariable=StringVar(),width=13,bg="white",font=style3)
    ad10a.grid(row=7,column=1,padx=10)
    ad11=Label(ad2,text="Residencial Address",width=20,anchor="w",bg=backcolour4,font=style3)
    ad11.grid(row=8,column=0,pady=2,padx=2)
    ad11a=Entry(ad2,textvariable=StringVar(),width=13,bg="white",font=style3)
    ad11a.grid(row=8,column=1,padx=10)
    ad12=Label(ad2,text="Permenant Address",width=20,anchor="w",bg=backcolour4,font=style3)
    ad12.grid(row=9,column=0,pady=2,padx=2)
    ad12a=Entry(ad2,textvariable=StringVar(),width=13,bg="white",font=style3)
    ad12a.grid(row=9,column=1,padx=10)
    pa=ad12a.get()
    ad13=Label(ad2,text="Common Rank",width=20,anchor="w",bg=backcolour4,font=style3)
    ad13.grid(row=10,column=0,pady=2,padx=2)
    ad13a=Entry(ad2,textvariable=IntVar(),width=13,bg="white",font=style3)
    ad13a.grid(row=10,column=1,padx=10)
    crl=ad13a.get()
    ad15=Label(ad2,text="Upload Document",width=20,anchor="w",bg=backcolour4,font=style3)
    ad15.grid(row=11,column=0,pady=4)
    ad16b=Button(ad2,text="Upload Aadhar",command=upload_file,anchor="w",width=15,font=style2)
    ad16b.grid(row=11,column=1,pady=3)
    ad16b=Button(ad2,text="Upload Migration",command=upload_file,anchor="w",width=15,font=style2)
    ad16b.grid(row=11,column=2,pady=3)
    ad16c=Button(ad2,text="Upload Rank Card",command=upload_file,anchor="w",width=15,font=style2)
    ad16c.grid(row=12,column=1,pady=3)
    ad16d=Button(ad2,text="Upload TC",command=upload_file,anchor="w",width=15,font=style2)
    ad16d.grid(row=12,column=2,pady=3)
    vary=StringVar()
    #TC,Aadhar,Migration,Rankcard
    #ad16a=Checkbutton(ad2,text="Aadhar",background=backcolour,fg="black",font=style2,activebackground=backcolour3,variable=vary,onvalue="NA",offvalue="Aadhar",command=filedialog.askopenfile(mode="r"))
    #ad16a.grid(row=12,column=0)
    ad14=Button(ad2,text="Submit and Generate Registration Number",bg="red",fg="white",command=generate_Admission_no)
    ad14.grid(row=13,column=0)
    Admission.bind('<Escape>',quit)
#============================================================
def student_details():
    mydb = mysql.connector.connect(
    host="localhost",
    user="root",
    password="Aditya@1234",
    database="Student_info")
    mycursor = mydb.cursor()
    detins= Val1_input.get()
    if detins=='':
        tmsg.showerror("Error","Unable To Fetch Data")
        return
    try:
        mycursor.execute(f"SELECT * FROM students where Registration={detins}")
        passkey = mycursor.fetchall()
        t1.config(state="normal")
        t1.delete(1.0,END)
        m=["Registartion:","Name:","Password:","Access:","Gender:","Branch:","Category:","Mother:","Father:","Residencial Address:","Permenant Address:","CRL:","Fees:","Fees Paid:","Fees Remaining:","DOB:","Mother Tounge:"]
        count=0
        for x in passkey:
            for y in x:
        #print(x)
                t1.insert(END,f"{m[count]}{y}\t")
                t1.insert(END,f"\n")
                count+=1
        t1.config(state="disabled")
        mycursor.close()
    except:
        tmsg.showerror("Error","Unable To Fetch Data")
#=============================================================
def capture():
    cam_port=0
    cam=cv2.VideoCapture(cam_port)
    width,height=100,100
    cam.set(cv2.CAP_PROP_FRAME_HEIGHT,height)
    cam.set(cv2.CAP_PROP_FRAME_WIDTH,width)
    ret,frame = cam.read()
    cam.release()
    if ret:
        global y
        cv2.imwrite(f"C:/Users/Adity/Desktop/mis/ID_Photoss/{inpidi.get()}.png",frame)
        y=f"C:/Users/Adity/Desktop/mis/ID_Photoss/{inpidi.get()}.png"
        upload_photo(y)
def upload_photo(photo_path):
    img = Image.open(photo_path)
    img = img.resize((200, 200), Image.LANCZOS) # Resize image if needed
    img = ImageTk.PhotoImage(img)
    label = Label(idg, image=img)
    label.image = img # Keep reference to avoid garbage collection
    label.place(relx=0.02,rely=0.28,relheight=0.3,relwidth=0.5)
def saveid():
    doc=docx.Document()
    doc.add_heading("\t\tABC INSTITUTE OF TECHNOLOGY",0)
    doc.add_picture(y,width=Inches(4))
    doc.add_paragraph(f"Employee ID:{inpidi.get()}")
    doc.add_paragraph(f"Name:{inpid1i.get()}")
    doc.add_paragraph(f"Department:{depm.get()}")
    doc.add_paragraph(f"Phone No:{inpid3i.get()}")
    #doc.add_paragraph(f"Date Of Joining:{Dt.get()}"),"Administrator","Visitor/Others"
    if tp.get()=="Student":
        doc.add_paragraph(f"Role:{tp.get()}")
        doc.add_paragraph(f"-------------------------------------------------------------------------------------------------------------------")
        doc.add_paragraph(f"\t\t\tABC Group Of Institute \n\t\t\t Greater Noida Delhi NCR INDIA")
        doc.save(f'C:/Users/Adity/Desktop/mis/ID_Inst/Student_ID/{inpidi.get()}.docx')
        tmsg.showinfo("Success","ID Card Generated")
    elif tp.get()=="Instructor":
        doc.add_paragraph(f"Role:{tp.get()}")
        doc.add_paragraph(f"-------------------------------------------------------------------------------------------------------------------")
        doc.add_paragraph(f"\t\t\tABC Group Of Institute \n\t\t\t Greater Noida Delhi NCR INDIA")
        doc.save(f'C:/Users/Adity/Desktop/mis/ID_Inst/Inst_ID/{inpidi.get()}.docx')
        tmsg.showinfo("Success","ID Card Generated")
    elif tp.get()=="Administrator":
        doc.add_paragraph(f"Role:{tp.get()}")
        doc.add_paragraph(f"-------------------------------------------------------------------------------------------------------------------")
        doc.add_paragraph(f"\t\t\tABC Group Of Institute \n\t\t\t Greater Noida Delhi NCR INDIA")
        doc.save(f'C:/Users/Adity/Desktop/mis/ID_Inst/Admin_ID/{inpidi.get()}.docx')
        tmsg.showinfo("Success","ID Card Generated")
    elif tp.get()=="Visitor/Others":
        doc.add_paragraph(f"Role:{tp.get()}")
        doc.add_paragraph(f"-------------------------------------------------------------------------------------------------------------------")
        doc.add_paragraph(f"\t\t\tABC Group Of Institute \n\t\t\t Greater Noida Delhi NCR INDIA")
        doc.save(f'C:/Users/Adity/Desktop/mis/ID_Inst/Misc/{inpidi.get()}.docx')
        tmsg.showinfo("Success","ID Card Generated for visitors")
    else:
        tmsg.showerror("Undefined","Select User Type")
def generate_id():
    global idg,inpidi,photolabel,inpid1i,depm,inpid3i,tp
    idgen=Toplevel(Win)
    idgen.geometry("430x480")
    #idgen.resizable(False,False)
    idgen.title("Id_Card_Generator")
    idg=Frame(idgen,bg="#31332f",relief="solid")
    idg.place(relx=0,rely=0,relheight=1,relwidth=1)
    inpidlabel=Label(idg,text="ID CARD GENERATION",font=style2,bg="grey",fg="white",relief="solid")
    inpidlabel.place(relheight=0.1,relwidth=1,relx=0,rely=0)
    inpid=Label(idg,text="Registration Number:",font=style2,bg="grey",fg="white",relief="raised")
    inpid.place(relx=0.02,rely=0.12,relheight=0.065,relwidth=0.5)
    inpidi=Entry(idg,textvariable=IntVar(),font=style2)
    inpidi.place(relx=0.54,rely=0.12,relheight=0.065,relwidth=0.4)
    inpidib=Button(idg,text="Click Photo",command=capture)
    inpidib.place(relx=0.02,rely=0.20,relheight=0.065,relwidth=0.5)
    photolabel=Label(idg)
    photolabel.place(relx=0.02,rely=0.28,relheight=0.3,relwidth=0.5)
    inpid1=Label(idg,text="Name",font=style2,bg="grey",fg="white",relief="raised")
    inpid1.place(relx=0.02,rely=0.61,relheight=0.065,relwidth=0.5)
    inpid1i=Entry(idg,textvariable=StringVar(),font=style2)
    inpid1i.place(relx=0.54,rely=0.61,relheight=0.065,relwidth=0.4)
    inpid2=Label(idg,text="Department",font=style2,bg="grey",fg="white",relief="raised")
    inpid2.place(relx=0.02,rely=0.71,relheight=0.065,relwidth=0.5)
    depm=StringVar()
    inpid2i=OptionMenu(idg,depm,"Computer Science","Information Technology","Civil Engineering","Mechanical Engineering","Chemical Engineering","Artificial Intelligence","Electronics and Communication","Mining Engineering")
    inpid2i.place(relx=0.54,rely=0.71,relheight=0.065,relwidth=0.4)
    depm.set("select")
    inpid3=Label(idg,text="Number",font=style2,bg="grey",fg="white",relief="raised")
    inpid3.place(relx=0.02,rely=0.81,relheight=0.065,relwidth=0.5)
    inpid3i=Entry(idg,textvariable=IntVar(),font=style2)
    inpid3i.place(relx=0.54,rely=0.81,relheight=0.065,relwidth=0.4)
    generid=Button(idg,text="Generate ID Card",font=style2,fg="white",bg="red",command=saveid)
    generid.place(relx=0.02,rely=0.91,relheight=0.065,relwidth=0.4)
    tp=StringVar()
    typeuser=OptionMenu(idg,tp,"Student","Instructor","Administrator","Visitor/Others")
    typeuser.place(relx=0,rely=0,relheight=0.05,relwidth=0.25)
    tp.set("Select")    
def up_onboard_photo():
    global full_path
    global img
    newpath = rf'C:\Users\Adity\Desktop\mis\Files\Instructor\{empidi.get()}' 
    if not os.path.exists(newpath):
        os.makedirs(newpath)
    file_path = filedialog.askopenfilename(initialdir='C:/Users/Adity/Downloads',filetypes=[("Image Files", "*.png *.jpg *.jpeg")],defaultextension='png')
    if file_path:
        full_path = os.path.abspath(file_path) 
        img = Image.open(file_path)
        img = img.resize((200, 200), Image.LANCZOS) # Resize image if needed
        img = ImageTk.PhotoImage(img)
        label = Label(onb, image=img)
        label.image = img # Keep reference to avoid garbage collection
        label.place(relx=0.7,rely=0.11,relheight=0.25,relwidth=0.25)
        saveupload(file_path,newpath)
def up_onboard_document():
    global full_path1
    global img
    newpath1 = rf'C:\Users\Adity\Desktop\mis\Files\Instructor\{empidi.get()}' 
    if not os.path.exists(newpath1):
        os.makedirs(newpath1)
    file_path1 = filedialog.askopenfilename(initialdir='C:/Users/Adity/Downloads',filetypes=[("Image Files", "*.png *.jpg *.jpeg")],defaultextension='png')
    saveupload(file_path1,newpath1)
def saveupload(file_path,newpath):
    save_path = filedialog.asksaveasfilename(initialdir=newpath,filetypes=[("Image Files", "*.png *.jpg *.jpeg")],title=f"Document{""}.jpg")
    if save_path:
        with open(file_path, "rb") as f:
            data = f.read()
        with open(save_path, "wb") as f:
            f.write(data)
    
#def upload_file():
 #   file_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg")])
  #  if file_path:
   #     os.mkdir()
def take_ob():
    doc=docx.Document()
    doc.add_heading("\t\tABC INSTITUTE OF TECHNOLOGY",0)
    doc.add_picture(full_path,width=Inches(4))
    doc.add_paragraph(f"Employee ID:{empidi.get()}")
    doc.add_paragraph(f"Name:{ot2i.get()}")
    doc.add_paragraph(f"Department:{depx.get()}")
    doc.add_paragraph(f"Phone No:{ot3i.get()}")
    doc.add_paragraph(f"Date Of Joining:{Dt.get()}")
    doc.add_paragraph(f"\t\t\tABC Group Of Institute \n\t\t\t Greater Noida Delhi NCR INDIA")
    doc.save(f'C:/Users/Adity/Desktop/mis/ID_Inst/{empidi.get()}.docx')
def fetchpayment():
    global d,val
    mydb = mysql.connector.connect(
  host="localhost",
  user="root",
  password="Aditya@1234",
  database="Teacher_info")
    mycursor = mydb.cursor()
    mycursor.execute(f"Select name,salary,AC_No from instructor where Emp_id={Emp2i.get()}")
    passkey=mycursor.fetchall()
    d=["Name:","Salary:","A/C No:"]
    val=[]
    count=0
    for i in passkey:
        x1.config(state=NORMAL)
        for j in i:
            x1.insert(END,f"{d[count]}\t{j}")
            x1.insert(END,f"\n")
            count+=1
            val.append(j)
    x1.config(state=DISABLED)
def salcred():
    Fet.config(state=DISABLED)
    tmsg.showinfo("Paid",f"Salary Credited To {val[2]}")
    x=tmsg.askyesnocancel("Payroll","Do you want to generate Payslip")
    if x==True:
        doc=docx.Document()
        doc.add_heading("\t\tABC INSTITUTE OF TECHNOLOGY",0)
        #doc.add_picture(full_path,width=Inches(4))
        doc.add_paragraph(f"Employee ID:{Emp2i.get()}")
        doc.add_paragraph(f"Name:{val[0]}")
        doc.add_paragraph(f"In Hand Payment:{int(val[1])/12}")
        doc.add_paragraph(f"Account Number:{val[2]}")
        doc.add_paragraph(f"Date Of Payment:{month.get_date()}")
        doc.add_paragraph(f"\t\t\tABC Group Of Institute \n\t\t\t Greater Noida Delhi NCR INDIA")
        d=str(month.get())
        e=d.replace("/","_")
        doc.save(f'C:/Users/Adity/Desktop/mis/Payslip/{Emp2i.get()}{e}.docx')
    

def paysalary():
    global payf1,Emp2i,x1,Fet,month
    pay=Toplevel(Win)
    pay.title("PAYMENT")
    pay.bell()
    pay.geometry("400x400")
    pay.maxsize(550,550)
    pay.minsize(350,350)
    payf1=Frame(pay,bg=backcolour)
    payf1.place(relx=0,rely=0,relwidth=1,relheight=1)
    Emp1=Label(payf1,text="PAYMENT",bg="grey",font=style2,relief=SUNKEN)
    Emp1.place(relheight=0.1,relwidth=1,relx=0,rely=0)
    Emp2=Label(payf1,text="Employee ID",bg="grey",font=style2,relief=SUNKEN)
    Emp2.place(relheight=0.085,relwidth=0.5,relx=0.02,rely=0.11)
    Emp2i=Entry(payf1,textvariable=IntVar(),font=style2)
    Emp2i.place(relx=0.54,rely=0.11,relheight=0.085,relwidth=0.4)
    Fet=Button(payf1,text="Fetch Details",bg="red",fg="white",command=fetchpayment)
    Fet.place(relx=0.02,rely=0.2,relheight=0.085,relwidth=0.5)
    x1=Text(payf1,bg=backcolour,border=0)
    x1.place(relx=0.02,rely=0.3,relwidth=1,relheight=0.15)
    pay=Button(payf1,text="Do you want to iniciate the payment?",command=salcred,bg="red")
    pay.place(relx=0.02,rely=0.47,relheight=0.085,relwidth=0.5)
    month=DateEntry(payf1)
    month.place(relx=0,rely=0)
#========================Take Onboard==========================
def on_board_data():
    valu=(empidi.get(),ot2i.get(),empidi.get(),0,ot5i.get(),str(ot3i.get()),ot4i.get(),ot6i.get(),ot7i.get(),depx.get(),qx.get(),sali.get(),str(Dt.get_date()),"null",anci.get())
    mydb = mysql.connector.connect(
    host="localhost",
    user="root",
    password="Aditya@1234",
    database="Teacher_info")
    mycursor = mydb.cursor()
    sql = f"INSERT INTO instructor (Emp_id,name,password,access,age,number,Gender,Email,Address,Dept,HighQ,salary,DOJ,DOE,AC_no) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
    mycursor.execute(sql,valu)
    mydb.commit()
    mydb.close()
    #"INSERT INTO instructor (Emp_id,name,password,access,age,number,Gender,Email,Address,Dept,HighQ,salary,DOJ,DOE,AC_no) VALUES
    tmsg.showinfo("Success",f"On board Complete Your \n ID: {empidi.get()}\n Password:{empidi.get()}")
    x=tmsg.askyesno("ID","Do you wan to generate ID Card?")
    if x==True:
        take_ob()
    else:
        onb.quit

    
def onboard():
    global ot2i,ot3i,empidi,ot5i,ot4i,ot6i,ot7i,depx,qx,sali,anci,Dt
    global onb
    global imglo
    onb=Toplevel(Win)
    onb.title("Onboard Client")
    onb.geometry("720x480")
    oF1=Frame(onb,bg="grey")
    oF1.place(relx=0,rely=0,relwidth=1,relheight=1)
    ot1=Label(oF1,text="Take Onboard",bg="lightgrey",fg="white",font=style)
    ot1.place(relx=0,rely=0,relwidth=1,relheight=0.07)
    ot2=Label(oF1,text="Name:",bg="grey",fg="white",font=style2)
    ot2.place(relx=0.02,rely=0.09,relheight=0.07,relwidth=0.25)
    ot3=Label(oF1,text="Number:",bg="grey",fg="white",font=style2)
    ot3.place(relx=0.02,rely=0.16,relheight=0.07,relwidth=0.25)
    ot4=Label(oF1,text="Gender(M/F):",bg="grey",fg="white",font=style2)
    ot4.place(relx=0.02,rely=0.23,relheight=0.07,relwidth=0.25)
    ot5=Label(oF1,text="Age:",bg="grey",fg="white",font=style2)
    ot5.place(relx=0.02,rely=0.3,relheight=0.07,relwidth=0.25)
    ot6=Label(oF1,text="Email:",bg="grey",fg="white",font=style2)
    ot6.place(relx=0.02,rely=0.37,relheight=0.07,relwidth=0.25)
    ot7=Label(oF1,text="Address:",bg="grey",fg="white",font=style2)
    ot7.place(relx=0.02,rely=0.44,relheight=0.14,relwidth=0.25)
    ot8=Label(oF1,text="Department:",bg="grey",fg="white",font=style2)
    ot8.place(relx=0.02,rely=0.58,relheight=0.07,relwidth=0.25)
    ot9=Label(oF1,text="Highest Qualification:",bg="grey",fg="white",font=style2)
    ot9.place(relx=0.02,rely=0.65,relheight=0.07,relwidth=0.25)
    ot10=Label(oF1,text="Upload Documents:",bg="grey",fg="white",font=style2)
    ot10.place(relx=0.02,rely=0.72,relheight=0.07,relwidth=0.25)
    ot2i=Entry(oF1,textvariable=StringVar(),font=style2)
    ot2i.place(relx=0.3,rely=0.09,relwidth=0.25,relheight=0.065)
    ot3i=Entry(oF1,textvariable=IntVar(),font=style2)
    ot3i.place(relx=0.3,rely=0.16,relwidth=0.25,relheight=0.065)
    ot4i=Entry(oF1,textvariable=StringVar(),font=style2)
    ot4i.place(relx=0.3,rely=0.23,relwidth=0.25,relheight=0.065)
    ot5i=Entry(oF1,textvariable=IntVar(),font=style2)
    ot5i.place(relx=0.3,rely=0.3,relwidth=0.25,relheight=0.065)
    ot6i=Entry(oF1,textvariable=StringVar(),font=style2)
    ot6i.place(relx=0.3,rely=0.37,relwidth=0.25,relheight=0.065)
    ot7i=Entry(oF1,textvariable=StringVar(),font=style2)
    ot7i.place(relx=0.3,rely=0.44,relwidth=0.25,relheight=0.130)
    depx=StringVar()
    ot8i=OptionMenu(oF1,depx,"Computer Science","Information Technology","Civil Engineering","Mechanical Engineering","Chemical Engineering","Artificial Intelligence","Electronics and Communication","Mining Engineering")
    depx.set("Select")
    ot8i.place(relx=0.3,rely=0.58,relwidth=0.25,relheight=0.065)
    qx=StringVar()
    ot9i=OptionMenu(oF1,qx,"Graduate","Post Graduate","Doctorate","Post Doctorate")
    qx.set("Select")
    ot9i.place(relx=0.3,rely=0.66,relwidth=0.25,relheight=0.065)
    ot10a=Button(oF1,text="Upload Photo",command=up_onboard_photo)
    ot10a.place(relx=0.02,rely=0.8,relwidth=0.25,relheight=0.065)
    ot10b=Button(oF1,text="Upload CV",command=up_onboard_document)
    ot10b.place(relx=0.3,rely=0.8,relwidth=0.25,relheight=0.065)
    ot10c=Button(oF1,text="Upload ID proof",command=up_onboard_document)
    ot10c.place(relx=0.02,rely=0.88,relwidth=0.25,relheight=0.065)
    ot10d=Button(oF1,text="Upload Degree & Experience",command=up_onboard_document)
    ot10d.place(relx=0.3,rely=0.88,relwidth=0.25,relheight=0.065)
    li=PhotoImage(height=20,width=20)
    imglo=Label(oF1,image=li)
    otph=Button(oF1,image=li,command=up_onboard_photo,borderwidth=0)
    otph.place(relx=0.7,rely=0.11,relheight=0.25,relwidth=0.25)  
    empid=Label(oF1,text="EmployeeId(Customised):",bg="grey",fg="white")
    empid.place(relx=0.7,rely=0.37,relwidth=0.25,relheight=0.065)
    empidi=Entry(oF1,textvariable=IntVar(),font=style2)
    empidi.place(relx=0.7,rely=0.45,relwidth=0.25,relheight=0.065)
    sal=Label(oF1,text="Salary:",bg="grey",fg="white",font=style2)
    sal.place(relx=0.7,rely=0.53,relheight=0.065,relwidth=0.25)
    sali=Entry(oF1,textvariable=IntVar(),font=style2)
    sali.place(relx=0.7,rely=0.59,relheight=0.065,relwidth=0.25)
    Dt=DateEntry(onb)
    Dt.place(relx=0,rely=0)
    Acn=Label(oF1,text="A/C No:",font=style2,bg="grey",fg="white")
    Acn.place(relx=0.7,rely=0.66,relheight=0.065,relwidth=0.25)
    anci=Entry(oF1,textvariable=IntVar(),font=style2)
    anci.place(relx=0.7,rely=0.73,relheight=0.065,relwidth=0.25)  
    ph=Label(otph,text="Photograph",fg="grey")
    ph.place(relx=0.35,rely=0.35)
    ot11=Button(oF1,text="Take On Board",bg="red",fg="white",command=on_board_data)

    ot11.place(relx=0.02,rely=0.95,relheight=0.05,relwidth=0.53)
def Ref():
    x=tmsg.showwarning("Relief","Do you want to Procced")
    if x==False:
        tmsg.showinfo("Action","Reverted")
    else:
        mydb = mysql.connector.connect(
        host="localhost",
        user="root",
        password="Aditya@1234",
        database="relief")
        mycursor = mydb.cursor()
        vl=(str(Emp_Idx.get()),str(Doe.get()),str(Roee.get()))
        mal=f"insert into relief values(%s,%s,%s)"
        mycursor.execute(mal,vl)
        mydb.commit()
        mydb.close()
        mydb = mysql.connector.connect(
        host="localhost",
        user="root",
        password="Aditya@1234",
        database="Teacher_info")
        mycursor = mydb.cursor()
        mycursor.execute(f"delete from instructor where Emp_id={(Emp_Idx.get())}")
        mydb.commit()
        mydb.close()
        tmsg.showinfo("Relief","Employee Relief Completed")
def push_announcment():
    mydb = mysql.connector.connect(
    host="localhost",
    user="root",
    password="Aditya@1234",
    database="announcment")
    mycursor = mydb.cursor()
    sql = f"INSERT INTO notice (note, target) VALUES (%s, %s)"   
    val = (af2.get("1.0",END),ty.get())
    mycursor.execute(sql, val)
    mydb.commit()
    mycursor.close()
    tmsg.showinfo("Push","Announcment Pushed")
def announcment():
    global ty,af2
    an=Toplevel(Win)
    an.title("Announcements")
    an.geometry("360x360")
    af=Frame(an,bg="grey")
    af.place(relx=0,rely=0,relheight=1,relwidth=1)
    af1=Label(af,text="Make Announements",font=style2,bg="grey")
    af1.place(relx=0,rely=0,relheight=0.085,relwidth=1)
    af2=Text(af,bg=backcolour,font=style2,fg="white")
    af2.place(relx=0,rely=0.09,relheight=0.75,relwidth=1)
    ty=StringVar(af)
    ty.set("Select")
    af4=OptionMenu(af,ty,"Students","Instructors","All")
    #ty.set("Select")
    af4.place(relx=0,rely=0.85,relheight=0.07,relwidth=0.4)  
    af3=Button(af,text="Push Announcements",bg="red",fg="white",command=push_announcment)
    af3.place(relx=0,rely=0.95,relheight=0.049,relwidth=0.4)
    
def feesdet():
    mydb = mysql.connector.connect(
    host="localhost",
    user="root",
    password="Aditya@1234",
    database="Student_info")
    mycursor = mydb.cursor()
    detins= Val1_input.get()
    if detins=='':
        tmsg.showerror("Error","Unable To Fetch Data")
        return
    try:
        mycursor.execute(f"SELECT fees,fees_paid,fees_remaining FROM students where Registration={detins}")
        passkey = mycursor.fetchall()
        t1.config(state="normal")
        t1.delete(1.0,END)
        m=["Fees:","Fees Paid:","Fees Remaining:"]
        count=0
        for x in passkey:
            for y in x:
        #print(x)
                t1.insert(END,f"{m[count]}{y}\t")
                t1.insert(END,f"\n")
                count+=1
        t1.config(state="disabled")
        mycursor.close()
    except:
        tmsg.showerror("Error","Unable To Fetch Data")
def Relief():
    global Emp_Idx,Doe,Roee
    Rs=Toplevel(Win)
    Rs.title("Relief")
    Rs.geometry("480x400")
    Rs.resizable(False,False)
    Frs=Frame(Rs,bg="grey")
    Frs.place(relx=0,rely=0,relwidth=1,relheight=1)
    Emp_Id=Label(Frs,text="Employee Id:",font=style,bg="grey",fg="white")
    Emp_Id.grid(column=0,row=0)
    Emp_Idx=Entry(Frs,font=style,textvariable=IntVar())
    Emp_Idx.grid(row=0,column=1)
    Doe1=Label(Frs,text="Date OF Relief:",font=style,bg="grey",fg="white")
    Doe1.grid(column=0,row=1)
    Doe=DateEntry(Frs)
    Doe.grid(row=1,column=1)
    Roe=Label(Frs,text="Reason OF Relief:",font=style,bg="grey",fg="white")
    Roe.grid(column=0,row=2)
    Roee=Entry(Frs,font=style,textvariable=StringVar())
    Roee.grid(row=2,column=1)
    acp=Checkbutton(Frs,bg="grey",fg="black",text="You hence certify that \n you have no objection",font=style2,activebackground="grey")
    acp.grid(row=3,column=0)
    cert=Button(Frs,text="Generate Certificate",command=generate_cert)
    cert.grid(row=4,column=0)
    Relieff=Button(Frs,text="Relief",command=Ref,bg="red",fg="white")
    Relieff.grid(row=5,column=0)
def Trf():
        if acpvar.get()==0:
            return
        else:
            mydb = mysql.connector.connect(
            host="localhost",
            user="root",
            password="Aditya@1234",
            database="transfer")
            mycursor = mydb.cursor()
            vl=(str(S_Idx.get()),str(Doeo.get()),str(Rope.get()),str(dpx.get()))
            mal=f"insert into transfer values(%s,%s,%s,%s)"
            mycursor.execute(mal,vl)
            mydb.commit()
            mydb.close()
            mydb = mysql.connector.connect(
            host="localhost",
            user="root",
            password="Aditya@1234",
            database="Student_Info")
            mycursor = mydb.cursor()
            mycursor.execute(f"delete from students where Registration={(S_Idx.get())}")
            mydb.commit()
            mydb.close()
            tmsg.showinfo("Relief","Student Transfer Completed")
    
def transfer():
    global S_Idx,Doeo,Rope,dpx,acpvar
    tf=Toplevel(Win)
    tf.title("Transfer")
    tf.geometry("480x480")
    Frs=Frame(tf,bg="grey")
    Frs.place(relx=0,rely=0,relwidth=1,relheight=1)
    S_Id=Label(Frs,text="Registration Id:",font=style,bg="grey",fg="white")
    S_Id.grid(column=0,row=0)
    S_Idx=Entry(Frs,font=style,textvariable=IntVar())
    S_Idx.grid(row=0,column=1)
    Doe1o=Label(Frs,text="Date OF Relief:",font=style,bg="grey",fg="white")
    Doe1o.grid(column=0,row=1)
    Doeo=DateEntry(Frs)
    Doeo.grid(row=1,column=1)
    Roeo=Label(Frs,text="Reason OF Relief:",font=style,bg="grey",fg="white")
    Roeo.grid(column=0,row=2)
    Rope=Entry(Frs,font=style,textvariable=StringVar())
    Rope.grid(row=2,column=1)
    dpx=StringVar()
    Dp=OptionMenu(Frs,dpx,"Computer Science","Information Technology","Civil Engineering","Mechanical Engineering","Chemical Engineering","Artificial Intelligence","Electronics and Communication","Mining Engineering")
    dpx.set("Select Department")
    Dp.grid(column=0,row=3)
    acpvar=IntVar()
    acp=Checkbutton(Frs,bg="grey",fg="black",text="You hence certify that \n you have no objection\n& no fees Due Remaining",font=style2,activebackground="grey",variable=acpvar,onvalue=1,offvalue=0)
    acp.grid(row=4,column=0)
    cert=Button(Frs,text="Generate Certificate",command=generate_cert)
    cert.grid(row=5,column=0)
    Relieff=Button(Frs,text="Transfer",command=Trf,bg="red",fg="white")
    Relieff.grid(row=6,column=0)



#===============================================================
Admin_Label=Label(Win,text="Administrator Panel",background=backcolour2,fg=front,font=style,relief=GROOVE)
Admin_Label.place(x=0,y=0,relwidth=1,height=50)
Back=Button(Win,text="<",font=style,command=back_command,background="black",fg="white")
Back.place(x=0,y=0,height=50,width=50)
F1=LabelFrame(Win,text="Client MIS Portal",bg=backcolour,font=style,labelanchor="n")
F1.place(x=2,y=55,relwidth=1,height=400)
Type=Radiobutton(F1,text="Student",font=style,background=backcolour,activebackground=backcolour,value="Student",variable=var,command=selected(var.get()),state="active")
Type.grid(row=0,column=0)
Type1=Radiobutton(F1,text="Instructor",font=style,background=backcolour,activebackground=backcolour,variable=var,value="Instructor",command=selected(var.get()))
Type1.grid(row=0,column=1)
Type2=Radiobutton(F1,text="Administrator",font=style,background=backcolour,activebackground=backcolour,variable=var,value="Administrator",command=selected(var.get()))
Type2.grid(row=0,column=2)
Val1=Label(F1,text="Registration no:",bg=backcolour,font=style)
Val1.grid(row=1,column=0)
Val1_input=Entry(F1,textvariable=IntVar,disabledbackground="white",background="white",font=style2)
Val1_input.grid(row=1,column=1,ipadx=4)
Val2=Label(F1,text="Name:",bg=backcolour,font=style)
Val2.grid(row=1,column=2)
Val2_input=Entry(F1,textvariable=StringVar,disabledbackground="white",background="white",font=style2)
Val2_input.grid(row=1,column=3,ipadx=0)
Val3=Label(F1,text="     Password:",bg=backcolour,font=style)
Val3.grid(row=1,column=4,padx=3)
Val3_input=Entry(F1,textvariable=StringVar,disabledbackground="white",background="white",font=style2)
Val3_input.grid(row=1,column=5,ipadx=4)
Val4=Checkbutton(F1,text="Do you want to make\n the changes applicable",background=backcolour,fg="black",font=style2,activebackground=backcolour,variable=var1,onvalue=1,offvalue=0)
Val4.grid(row=3,column=0,pady=10)
Val5=Button(F1,text="Add User",background=Buttoncolor,font=style,fg="white",activebackground=backcolour,width=15,command=add_user)
Val5.grid(row=5,column=0,padx=15)
Val6=Button(F1,text="Change User Data",background=Buttoncolor,fg="white",font=style,activebackground=backcolour,width=15,command=change_data)
Val6.grid(row=5,column=1,padx=15)
Val7=Button(F1,text="Revoke User",background=Buttoncolor,fg="white",font=style,activebackground=backcolour,width=15,command=revoke_user)
Val7.grid(row=5,column=2,padx=15)
Val8=Label(F1,text=captcha,font=style3,bg="#dbcad0",width=15)
Val8.grid(row=4,column=0,pady=10)
Val8_input=Entry(F1,textvariable=StringVar,background="white",font=style3,width=15)
Val8_input.grid(row=4,column=1,padx=10,pady=10)
Val9=Button(F1,text="Allow User",background=Buttoncolor,fg="white",font=style,activebackground=backcolour,width=15,command=allow_access)
Val9.grid(row=6,column=0,pady=10)
Val10=Button(F1,text="Reset",background=Buttoncolor,fg="white",font=style,activebackground=backcolour,width=15,command=reset)
Val10.grid(row=6,column=1,pady=10)
Val11=Button(F1,text="Delete",background=Buttoncolor,fg="white",font=style,activebackground=backcolour,width=15,command=delete_user)
Val11.grid(row=6,column=2,pady=10)
F2=LabelFrame(Win,text="Details",font=style,labelanchor="n")
F2.place(x=5,y=460,relwidth=0.5,relheight=0.43)
t1=Text(F2,state="normal",padx=1,pady=1,font=style3)
t1.place(x=0,y=0,relwidth=1,relheight=0.75)
scrollbar = Scrollbar(t1)
scrollbar.pack( side = RIGHT, fill=Y )
scrollbar.config(command=t1.yview,jump=1,orient=VERTICAL)
t1.config(yscrollcommand=scrollbar.set)
ValF2_1=Button(F2,text="Fetch all data",background=Buttoncolor,fg="white",font=style,activebackground=backcolour,width=15,command=get_data)
ValF2_1.place(x=50,y=260)
ValF2_2=Button(F2,text="Reset",background=Buttoncolor,fg="white",font=style,activebackground=backcolour,width=15,command=reset_1)
ValF2_2.place(x=400,y=260)
F3=LabelFrame(Win,text="Action Pane",labelanchor="n",font=style)
F3.place(x=800,y=460,relheight=0.43,relwidth=0.47)
Actp=Button(F3,text="Get Complete Student Details",height=3,width=30,command=student_details)
Actp.grid(row=0,column=0,padx=4,pady=4)
Actp1=Button(F3,text="Get Student Fees Details",height=3,width=30,command=feesdet)
Actp1.grid(row=1,column=0,padx=4,pady=4)
Actp2=Button(F3,text="See Complete Student Document",height=3,width=30,command=stu_files)
Actp2.grid(row=2,column=0,padx=4,pady=4)
Actp3=Button(F3,text="Make new Admission",height=3,width=30,command=make_admission)
Actp3.grid(row=3,column=0,padx=4,pady=4)
Actp4=Button(F3,text="Transfer Student",height=3,width=30,command=transfer)
Actp4.grid(row=4,column=0,padx=4,pady=4)
#============================================================#
EActp=Button(F3,text="Get Complete Employee Details",height=3,width=30,command=show_det_inst)
EActp.grid(row=0,column=1,padx=4,pady=4)
EActp1=Button(F3,text="Pay Employee Salary",height=3,width=30,command=paysalary)
EActp1.grid(row=1,column=1,padx=4,pady=4)
EActp2=Button(F3,text="See Employee Document",height=3,width=30,command=instruct_files)
EActp2.grid(row=2,column=1,padx=4,pady=4)
EActp3=Button(F3,text="Take Onboard",height=3,width=30,command=onboard)
EActp3.grid(row=3,column=1,padx=4,pady=4)
EActp4=Button(F3,text="Relief Instructor",height=3,width=30,command=Relief)
EActp4.grid(row=4,column=1,padx=4,pady=4)
#================================================================
EActp=Button(F3,text="Push Alerts",height=3,width=30,command=announcment)
EActp.grid(row=0,column=2,padx=4,pady=4)
EActp1=Button(F3,text="Generate Certificates",height=3,width=30,command=generate_cert)
EActp1.grid(row=1,column=2,padx=4,pady=4)
EActp2=Button(F3,text="Generate ID Cards",height=3,width=30,command=generate_id)
EActp2.grid(row=2,column=2,padx=4,pady=4)
EActp3=Button(F3,text="Check Student Attendence",height=3,width=30)
EActp3.grid(row=3,column=2,padx=4,pady=4)
EActp4=Button(F3,text="Instructor Attendence",height=3,width=30)
EActp4.grid(row=4,column=2,padx=4,pady=4)
Copy=Label(Win,text="COPYRIGHTS ABC INSTITUTE OF TECHNOLOGY",font=("timesnewroman",10,"bold"),border=2,relief="raise")
Copy.pack(fill=X,anchor="s",side="bottom")
#===================Closing Attributes==========================
Win.attributes("-fullscreen",True)
Win.bind('<Escape>',quit)
Win.mainloop()
